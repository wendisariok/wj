"""Word document generation for email export and collection book export."""

import io
from datetime import datetime
from typing import Optional

import html2text
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _html_to_text(html: Optional[str]) -> str:
    """Convert HTML to plain text."""
    if not html:
        return ""
    h = html2text.HTML2Text()
    h.ignore_links = False
    h.ignore_images = True
    h.body_width = 0
    return h.handle(html).strip()


def _get_email_body(email: dict) -> str:
    """Get the best available body text from an email."""
    if email.get("body_text"):
        return email["body_text"]
    if email.get("body_html"):
        return _html_to_text(email["body_html"])
    return email.get("snippet") or "(no content)"


def _format_date(date_str: Optional[str]) -> str:
    """Format an ISO date string for display."""
    if not date_str:
        return "Unknown date"
    try:
        dt = datetime.fromisoformat(date_str)
        return dt.strftime("%B %d, %Y %I:%M %p")
    except (ValueError, TypeError):
        return date_str


def _get_month_key(date_str: Optional[str]) -> str:
    """Extract month/year key for grouping."""
    if not date_str:
        return "Unknown Date"
    try:
        dt = datetime.fromisoformat(date_str)
        return dt.strftime("%B %Y")
    except (ValueError, TypeError):
        return "Unknown Date"


def _add_title_page(doc: Document, title: str, subtitle: str = ""):
    """Add a title page to the document."""
    doc.add_paragraph()
    doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    if subtitle:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub_para.add_run(subtitle)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"Exported {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()


def _format_file_size(size: Optional[int]) -> str:
    """Format a file size in bytes for display."""
    if size is None or size == 0:
        return ""
    if size < 1024:
        return f"{size} B"
    if size < 1024 * 1024:
        return f"{size / 1024:.1f} KB"
    return f"{size / (1024 * 1024):.1f} MB"


def _add_email_to_doc(doc: Document, email: dict, attachments: Optional[list[dict]] = None):
    """Add a single email to the document."""
    # Subject heading
    subject = email.get("subject") or "(no subject)"
    heading = doc.add_heading(subject, level=2)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    # Metadata table
    table = doc.add_table(rows=3, cols=2)
    table.style = "Light List"

    cells = table.rows[0].cells
    cells[0].text = "From"
    cells[1].text = email.get("sender") or "Unknown"

    cells = table.rows[1].cells
    cells[0].text = "To"
    cells[1].text = email.get("recipient") or "Unknown"

    cells = table.rows[2].cells
    cells[0].text = "Date"
    cells[1].text = _format_date(email.get("date"))

    doc.add_paragraph()

    # Body
    body = _get_email_body(email)
    for line in body.split("\n"):
        doc.add_paragraph(line)

    # Attachments list
    if attachments:
        att_para = doc.add_paragraph()
        run = att_para.add_run("Attachments:")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        for att in attachments:
            filename = att.get("filename") or "unnamed"
            size_str = _format_file_size(att.get("size"))
            label = f"  {filename}" + (f" ({size_str})" if size_str else "")
            p = doc.add_paragraph(label)
            for r in p.runs:
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()


def build_emails_docx(
    emails: list[dict],
    group_by: str = "none",
    title: str = "Email Export",
    attachments_by_email: Optional[dict[int, list[dict]]] = None,
) -> bytes:
    """Build a Word document from a list of emails.

    Args:
        emails: List of email dicts (from DB rows).
        group_by: "none", "date", or "sender".
        title: Document title for the title page.
        attachments_by_email: Optional mapping of email_id -> list of attachment dicts.

    Returns:
        Bytes of the .docx file.
    """
    doc = Document()
    att_map = attachments_by_email or {}

    _add_title_page(doc, title, f"{len(emails)} email(s)")

    if group_by == "date":
        groups: dict[str, list[dict]] = {}
        for email in emails:
            key = _get_month_key(email.get("date"))
            groups.setdefault(key, []).append(email)
        # Sort reverse chronological
        for group_name in sorted(groups.keys(), reverse=True):
            doc.add_heading(group_name, level=1)
            for email in groups[group_name]:
                _add_email_to_doc(doc, email, att_map.get(email.get("id")))

    elif group_by == "sender":
        groups = {}
        for email in emails:
            key = email.get("sender") or "Unknown Sender"
            groups.setdefault(key, []).append(email)
        for group_name in sorted(groups.keys()):
            doc.add_heading(group_name, level=1)
            for email in groups[group_name]:
                _add_email_to_doc(doc, email, att_map.get(email.get("id")))

    else:
        for email in emails:
            _add_email_to_doc(doc, email, att_map.get(email.get("id")))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_collection_docx(
    title: str,
    description: str,
    entries: list[dict],
    attachments_by_email: Optional[dict[int, list[dict]]] = None,
) -> bytes:
    """Build a structured book document from a collection of emails.

    Args:
        title: Book/collection title.
        description: Collection description.
        entries: List of dicts with 'chapter_title', 'subject', 'sender',
                 'recipient', 'date', 'body_text', 'body_html', 'snippet'.
        attachments_by_email: Optional mapping of email_id -> list of attachment dicts.

    Returns:
        Bytes of the .docx file.
    """
    doc = Document()
    att_map = attachments_by_email or {}

    _add_title_page(doc, title, description)

    # Table of contents
    if entries:
        toc_heading = doc.add_heading("Table of Contents", level=1)
        for run in toc_heading.runs:
            run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

        for i, entry in enumerate(entries, 1):
            chapter = entry.get("chapter_title") or entry.get("subject") or f"Chapter {i}"
            toc_para = doc.add_paragraph()
            run = toc_para.add_run(f"{i}. {chapter}")
            run.font.size = Pt(11)

        doc.add_page_break()

    # Chapters
    for i, entry in enumerate(entries, 1):
        chapter = entry.get("chapter_title") or entry.get("subject") or f"Chapter {i}"
        doc.add_heading(chapter, level=1)
        email_id = entry.get("id") or entry.get("email_id")
        _add_email_to_doc(doc, entry, att_map.get(email_id))
        if i < len(entries):
            doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
